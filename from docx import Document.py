from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load the document
doc = Document("C:/Users/sgrka/OneDrive/Desktop/HeartGuardProject-Assignment4.docx")

# Set font to Calibri and size to 12
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

# Set line spacing to 1.5
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5

# Set margins to Normal (1 inch = 2.54 cm)
sections = doc.sections
for section in sections:
    section.top_margin = Pt(72)  # 1 inch
    section.bottom_margin = Pt(72)  # 1 inch
    section.left_margin = Pt(72)  # 1 inch
    section.right_margin = Pt(72)  # 1 inch

# Format Title
title_texts = [
    "HUMBER INSTITUTE OF TECHNOLOGY",
    "AND ADVANCED LEARNING",
    "(HUMBER COLLEGE)",
    "",
    "TEAM 3",
    "ASSIGNMENT 4",
    "CAPSTONE PROJECT {BIA-5450-0LA}",
    "HEARTGUARD PROJECT",
]

# Center align and bold the title texts
for paragraph in doc.paragraphs[:8]:
    if paragraph.text.strip() in title_texts:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.bold = True

# Format table
table = doc.tables[0]
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Ensure each sentence starts a new paragraph
new_paragraphs = []
for paragraph in doc.paragraphs[8:]:
    sentences = paragraph.text.split('. ')
    for sentence in sentences:
        new_paragraph = OxmlElement("w:p")
        new_run = OxmlElement("w:r")
        new_text = OxmlElement("w:t")
        new_text.text = sentence.strip() + ('.' if sentence else '')
        new_run.append(new_text)
        new_paragraph.append(new_run)
        new_paragraphs.append(new_paragraph)

body = doc._element.body
for paragraph in new_paragraphs:
    body.append(paragraph)

# Save the updated document
updated_file_path = '/mnt/data/HeartGuardProject-Assignment4-Formatted.docx'
doc.save(updated_file_path)

updated_file_path
